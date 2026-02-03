from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import sys

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def convert_bab3_to_docx(input_file, output_file):
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].rstrip('\r')
        
        # BAB heading
        if line.startswith('BAB III'):
            p = doc.add_heading(line, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # Sub-section headings (lines that are titles without indentation)
        if line and not line.startswith('\t') and not line.startswith(' ') and not line.startswith('│') and not line.startswith('├') and not line.startswith('└') and not line.startswith('┌') and not line.startswith('─') and not line.startswith('$') and not line.startswith('Gambar') and not line.startswith('Tabel') and not line.startswith('http') and not line.startswith('URL'):
            # Check if it's a section title (no special chars, reasonable length)
            if len(line) < 80 and not any(c in line for c in ['│', '├', '└', '┌', '─', '$', '>', '=']):
                # Check if next lines are content (not empty or table)
                if i + 1 < len(lines):
                    next_line = lines[i + 1].rstrip('\r') if i + 1 < len(lines) else ''
                    if next_line.startswith('\t') or next_line == '' or next_line.startswith('┌'):
                        # This is a section heading
                        if any(keyword in line for keyword in ['Persiapan', 'Pembuatan', 'Konfigurasi', 'Akses', 'Instalasi', 'Clone', 'Import', 'Testing', 'Troubleshooting', 'Kesimpulan', 'Pemilihan']):
                            doc.add_heading(line, level=1)
                            i += 1
                            continue
                        elif any(keyword in line for keyword in ['Update', 'Nginx', 'PHP', 'MySQL', 'Composer', 'Node', 'Virtual', 'Aktivasi', 'Repository', 'Permission', 'Dependencies', 'Build', 'Membuat', 'Export', 'Upload', 'Setup', 'Generate', 'Test', 'Akses Website']):
                            doc.add_heading(line, level=2)
                            i += 1
                            continue
        
        # Gambar placeholder
        if line.startswith('Gambar'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.italic = True
            run.font.size = Pt(10)
            i += 1
            continue
        
        # Table detection - look for table header (Tabel X.X)
        if line.startswith('Tabel'):
            table_title = line
            i += 1
            
            # Collect table content
            table_content = []
            while i < len(lines) and (lines[i].rstrip('\r').startswith('┌') or 
                                       lines[i].rstrip('\r').startswith('│') or 
                                       lines[i].rstrip('\r').startswith('├') or 
                                       lines[i].rstrip('\r').startswith('└') or
                                       lines[i].rstrip('\r').startswith('─')):
                table_content.append(lines[i].rstrip('\r'))
                i += 1
            
            if table_content:
                # Add table title
                p = doc.add_paragraph()
                run = p.add_run(table_title)
                run.bold = True
                run.font.size = Pt(11)
                
                # Create a single-column table for code/command boxes
                tbl = doc.add_table(rows=1, cols=1)
                tbl.style = 'Table Grid'
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Set table width
                tbl.columns[0].width = Inches(6.5)
                
                # Extract content from box
                content_lines = []
                for tl in table_content:
                    # Remove box characters and extract text
                    cleaned = tl.replace('┌', '').replace('┐', '').replace('└', '').replace('┘', '')
                    cleaned = cleaned.replace('├', '').replace('┤', '').replace('│', '')
                    cleaned = cleaned.replace('─', '').strip()
                    if cleaned:
                        content_lines.append(cleaned)
                
                cell = tbl.rows[0].cells[0]
                # Set light gray background for code blocks
                set_cell_shading(cell, 'F5F5F5')
                
                cell_para = cell.paragraphs[0]
                for idx, cl in enumerate(content_lines):
                    if idx > 0:
                        cell_para.add_run('\n')
                    run = cell_para.add_run(cl)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                
                doc.add_paragraph()  # Add spacing after table
            continue
        
        # Regular paragraph with tab indentation
        if line.startswith('\t'):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            p.add_run(line.strip())
            i += 1
            continue
        
        # URL line
        if line.startswith('URL:') or line.startswith('http'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 102, 204)
            i += 1
            continue
        
        # Regular text (bullet points or normal)
        if line.strip():
            if line.strip().startswith('-') or line.strip().startswith('•'):
                p = doc.add_paragraph(line.strip()[1:].strip(), style='List Bullet')
            elif re.match(r'^\d+\.', line.strip()):
                p = doc.add_paragraph(re.sub(r'^\d+\.\s*', '', line.strip()), style='List Number')
            else:
                p = doc.add_paragraph(line.strip())
            i += 1
            continue
        
        i += 1
    
    doc.save(output_file)
    print(f'Successfully saved to {output_file}')

if __name__ == '__main__':
    if len(sys.argv) >= 3:
        convert_bab3_to_docx(sys.argv[1], sys.argv[2])
    else:
        convert_bab3_to_docx('BAB_III_VPS.txt', 'BAB_III_VPS.docx')

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import sys

def md_to_docx(md_file, docx_file):
    doc = Document()
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    in_code_block = False
    code_content = []
    
    for line in lines:
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_content)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_content.append(line)
            continue
        
        # Handle headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)
        # Handle horizontal rules
        elif line.startswith('---'):
            doc.add_paragraph('─' * 50)
        # Handle bullet points
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\. ', line):
            doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
        # Handle blockquotes
        elif line.startswith('> '):
            p = doc.add_paragraph(line[2:])
            p.paragraph_format.left_indent = Inches(0.5)
        # Handle tables (simplified)
        elif line.startswith('|'):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if not all(c.replace('-', '') == '' for c in cells):
                p = doc.add_paragraph(' | '.join(cells))
                p.runs[0].font.size = Pt(10)
        # Regular paragraph
        elif line.strip():
            # Clean up markdown formatting
            clean_line = line.replace('**', '').replace('`', '').replace('*', '')
            doc.add_paragraph(clean_line)
    
    doc.save(docx_file)
    print(f'Saved to {docx_file}')

if __name__ == '__main__':
    if len(sys.argv) >= 3:
        md_to_docx(sys.argv[1], sys.argv[2])
    else:
        print('Usage: python md_to_docx.py input.md output.docx')
